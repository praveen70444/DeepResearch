"""Document ingestion and content extraction."""

import os
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional
import logging

# Document processing imports
import PyPDF2
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import markdown

from ..models import Document, TextChunk, DocumentFormat
from ..interfaces import DocumentIngesterInterface
from ..exceptions import DocumentIngestionError
from ..config import config

logger = logging.getLogger(__name__)


class DocumentIngester(DocumentIngesterInterface):
    """Handles document ingestion and content extraction."""
    
    def __init__(self):
        """Initialize document ingester."""
        self._supported_formats = [
            DocumentFormat.PDF,
            DocumentFormat.TXT,
            DocumentFormat.DOCX,
            DocumentFormat.MD,
            DocumentFormat.HTML
        ]
    
    def supported_formats(self) -> List[str]:
        """Return list of supported file formats."""
        return [fmt.value for fmt in self._supported_formats]
    
    def ingest_document(self, file_path: str) -> List[TextChunk]:
        """
        Ingest a document and return text chunks.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of TextChunk objects
            
        Raises:
            DocumentIngestionError: If ingestion fails
        """
        try:
            # Validate file exists
            if not os.path.exists(file_path):
                raise DocumentIngestionError(f"File not found: {file_path}")
            
            # Detect format
            format_type = self._detect_format(file_path)
            if format_type not in self._supported_formats:
                raise DocumentIngestionError(f"Unsupported format: {format_type}")
            
            # Extract content
            content, metadata = self._extract_content(file_path, format_type)
            
            if not content.strip():
                raise DocumentIngestionError(f"No content extracted from: {file_path}")
            
            # Create document
            document = Document(
                id=str(uuid.uuid4()),
                title=self._extract_title(file_path, content, metadata),
                content=content,
                source_path=file_path,
                format_type=format_type,
                metadata=metadata
            )
            
            # Create chunks (will be handled by TextProcessor in next task)
            # For now, return empty list - chunks will be created by TextProcessor
            logger.info(f"Successfully ingested document: {file_path}")
            return []
            
        except Exception as e:
            logger.error(f"Failed to ingest document {file_path}: {e}")
            raise DocumentIngestionError(f"Ingestion failed for {file_path}: {e}")
    
    def ingest_document_full(self, file_path: str) -> Document:
        """
        Ingest a document and return the full Document object.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Document object with content and metadata
        """
        try:
            # Validate file exists
            if not os.path.exists(file_path):
                raise DocumentIngestionError(f"File not found: {file_path}")
            
            # Detect format
            format_type = self._detect_format(file_path)
            if format_type not in self._supported_formats:
                raise DocumentIngestionError(f"Unsupported format: {format_type}")
            
            # Extract content
            content, metadata = self._extract_content(file_path, format_type)
            
            if not content.strip():
                raise DocumentIngestionError(f"No content extracted from: {file_path}")
            
            # Create document
            document = Document(
                id=str(uuid.uuid4()),
                title=self._extract_title(file_path, content, metadata),
                content=content,
                source_path=file_path,
                format_type=format_type,
                metadata=metadata
            )
            
            logger.info(f"Successfully ingested document: {file_path}")
            return document
            
        except Exception as e:
            logger.error(f"Failed to ingest document {file_path}: {e}")
            raise DocumentIngestionError(f"Ingestion failed for {file_path}: {e}")
    
    def _detect_format(self, file_path: str) -> DocumentFormat:
        """Detect document format from file extension."""
        extension = Path(file_path).suffix.lower()
        
        format_map = {
            '.pdf': DocumentFormat.PDF,
            '.txt': DocumentFormat.TXT,
            '.docx': DocumentFormat.DOCX,
            '.doc': DocumentFormat.DOCX,  # Treat .doc as .docx
            '.md': DocumentFormat.MD,
            '.markdown': DocumentFormat.MD,
            '.html': DocumentFormat.HTML,
            '.htm': DocumentFormat.HTML
        }
        
        if extension not in format_map:
            raise DocumentIngestionError(f"Unsupported file extension: {extension}")
        
        return format_map[extension]
    
    def _extract_content(self, file_path: str, format_type: DocumentFormat) -> tuple[str, Dict[str, Any]]:
        """
        Extract content and metadata from document.
        
        Returns:
            Tuple of (content, metadata)
        """
        try:
            if format_type == DocumentFormat.PDF:
                return self._extract_pdf_content(file_path)
            elif format_type == DocumentFormat.TXT:
                return self._extract_txt_content(file_path)
            elif format_type == DocumentFormat.DOCX:
                return self._extract_docx_content(file_path)
            elif format_type == DocumentFormat.MD:
                return self._extract_markdown_content(file_path)
            elif format_type == DocumentFormat.HTML:
                return self._extract_html_content(file_path)
            else:
                raise DocumentIngestionError(f"No extractor for format: {format_type}")
                
        except Exception as e:
            raise DocumentIngestionError(f"Content extraction failed: {e}")
    
    def _extract_pdf_content(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract content from PDF file."""
        content = ""
        metadata = {}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata.update({
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'producer': pdf_reader.metadata.get('/Producer', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                        'modification_date': str(pdf_reader.metadata.get('/ModDate', ''))
                    })
                
                metadata['page_count'] = len(pdf_reader.pages)
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            content += f"\n--- Page {page_num + 1} ---\n"
                            content += page_text
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                        continue
                
        except Exception as e:
            raise DocumentIngestionError(f"PDF extraction failed: {e}")
        
        return content.strip(), metadata
    
    def _extract_txt_content(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract content from text file."""
        metadata = {}
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        metadata['encoding'] = encoding
                        break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise DocumentIngestionError("Could not decode text file with any supported encoding")
            
            # Add file stats
            file_stat = os.stat(file_path)
            metadata.update({
                'file_size': file_stat.st_size,
                'line_count': len(content.splitlines())
            })
            
        except Exception as e:
            raise DocumentIngestionError(f"Text extraction failed: {e}")
        
        return content, metadata
    
    def _extract_docx_content(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract content from DOCX file."""
        content = ""
        metadata = {}
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata.update({
                'author': core_props.author or '',
                'title': core_props.title or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
                'revision': core_props.revision or 0
            })
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
            
            content = '\n\n'.join(paragraphs)
            metadata['paragraph_count'] = len(paragraphs)
            
        except Exception as e:
            raise DocumentIngestionError(f"DOCX extraction failed: {e}")
        
        return content, metadata
    
    def _extract_markdown_content(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract content from Markdown file."""
        metadata = {}
        
        try:
            # Read raw markdown
            with open(file_path, 'r', encoding='utf-8') as file:
                raw_content = file.read()
            
            # Convert to HTML then extract text
            html_content = markdown.markdown(raw_content)
            soup = BeautifulSoup(html_content, 'html.parser')
            content = soup.get_text()
            
            # Extract metadata from front matter if present
            if raw_content.startswith('---'):
                parts = raw_content.split('---', 2)
                if len(parts) >= 3:
                    # Simple front matter parsing (YAML-like)
                    front_matter = parts[1].strip()
                    for line in front_matter.split('\n'):
                        if ':' in line:
                            key, value = line.split(':', 1)
                            metadata[key.strip()] = value.strip()
            
            metadata.update({
                'raw_length': len(raw_content),
                'processed_length': len(content),
                'has_front_matter': raw_content.startswith('---')
            })
            
        except Exception as e:
            raise DocumentIngestionError(f"Markdown extraction failed: {e}")
        
        return content, metadata
    
    def _extract_html_content(self, file_path: str) -> tuple[str, Dict[str, Any]]:
        """Extract content from HTML file."""
        metadata = {}
        
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract metadata from HTML head
            if soup.title:
                metadata['html_title'] = soup.title.string
            
            # Extract meta tags
            meta_tags = soup.find_all('meta')
            for meta in meta_tags:
                if meta.get('name') and meta.get('content'):
                    metadata[f"meta_{meta.get('name')}"] = meta.get('content')
            
            # Extract text content
            content = soup.get_text()
            
            # Clean up whitespace
            lines = [line.strip() for line in content.splitlines()]
            content = '\n'.join(line for line in lines if line)
            
            metadata.update({
                'html_length': len(html_content),
                'text_length': len(content),
                'has_title': bool(soup.title)
            })
            
        except Exception as e:
            raise DocumentIngestionError(f"HTML extraction failed: {e}")
        
        return content, metadata
    
    def _extract_title(self, file_path: str, content: str, metadata: Dict[str, Any]) -> str:
        """Extract or generate document title."""
        # Try to get title from metadata first
        title_candidates = [
            metadata.get('title'),
            metadata.get('html_title'),
            metadata.get('/Title')  # PDF title
        ]
        
        for candidate in title_candidates:
            if candidate and candidate.strip():
                return candidate.strip()
        
        # Try to extract from first line of content
        first_line = content.split('\n')[0].strip()
        if first_line and len(first_line) < 200:  # Reasonable title length
            return first_line
        
        # Fall back to filename
        return Path(file_path).stem
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic file information without full ingestion."""
        try:
            if not os.path.exists(file_path):
                raise DocumentIngestionError(f"File not found: {file_path}")
            
            file_stat = os.stat(file_path)
            format_type = self._detect_format(file_path)
            
            return {
                'file_path': file_path,
                'file_name': Path(file_path).name,
                'file_size': file_stat.st_size,
                'format_type': format_type.value,
                'modified_time': file_stat.st_mtime,
                'is_supported': format_type in self._supported_formats
            }
            
        except Exception as e:
            raise DocumentIngestionError(f"Failed to get file info: {e}")
    
    def batch_ingest(self, file_paths: List[str]) -> List[Document]:
        """
        Ingest multiple documents in batch.
        
        Args:
            file_paths: List of file paths to ingest
            
        Returns:
            List of successfully ingested documents
        """
        documents = []
        errors = []
        
        for file_path in file_paths:
            try:
                document = self.ingest_document_full(file_path)
                documents.append(document)
            except DocumentIngestionError as e:
                errors.append(f"{file_path}: {e}")
                logger.error(f"Batch ingestion error: {e}")
        
        if errors:
            logger.warning(f"Batch ingestion completed with {len(errors)} errors")
        
        return documents